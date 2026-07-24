from fastapi import APIRouter
from fastapi.responses import FileResponse
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os

router = APIRouter()


@router.post("/export/pdf")
async def export_pdf(data: dict):

    os.makedirs("generated_reports", exist_ok=True)

    filename = "generated_reports/summary.pdf"

    doc = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()

    story = [
        Paragraph("<b>AI Summary</b>", styles["Heading1"]),
        Paragraph(data["content"], styles["BodyText"]),
    ]

    doc.build(story)

    return FileResponse(
        filename,
        filename="Summary.pdf"
    )


@router.post("/export/docx")
async def export_docx(data: dict):

    os.makedirs("generated_reports", exist_ok=True)

    filename = "generated_reports/summary.docx"

    doc = Document()

    doc.add_heading("AI Summary", level=1)

    doc.add_paragraph(data["content"])

    doc.save(filename)

    return FileResponse(
        filename,
        filename="Summary.docx"
    )